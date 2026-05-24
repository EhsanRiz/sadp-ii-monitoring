"""Parse ESMP .docx files into form-schema JSON."""
import os, json, re, sys
from docx import Document

BASE = "/sessions/brave-compassionate-ramanujan/mnt/ESMP forms"
OUT  = "/tmp/esmp_schemas.json"

CHECKBOX_LEAD = re.compile(r'^[\s☐☑☒□■✓✔●•·\-\*]+')
SECTION_RE   = re.compile(r'^(\d+)\.0\b')          # 1.0, 2.0, 3.0 — section header
ACTIVITY_RE  = re.compile(r'^(\d+\.\d+)\b')        # 1.1, 2.1, 3.4 — sub-activity

def cell_items(cell):
    items = []
    for p in cell.paragraphs:
        t = p.text.strip()
        if not t: continue
        t = CHECKBOX_LEAD.sub('', t).strip()
        if t: items.append(t)
    return items

def cell_plain(cell):
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return ' '.join(parts) if parts else None

def detect_section(row):
    """A section header row: either fully-merged with same text, or col 0 starts
    with N.0 (and other cells empty or different)."""
    c0 = row.cells[0].text.strip()
    if SECTION_RE.match(c0):
        return c0.replace('\n', ' ').strip()
    return None

def parse_emmp(path):
    doc = Document(path)
    title_paragraphs = [p.text.strip() for p in doc.paragraphs[:5] if p.text.strip()]
    title = title_paragraphs[0] if title_paragraphs else os.path.basename(path)
    version_match = re.search(r'(Ver\.?\s*\d[^\)]*)', os.path.basename(path))
    version = version_match.group(1).strip() if version_match else 'Ver. 1'

    if not doc.tables:
        return None
    table = doc.tables[0]
    sections = []
    current_section = None

    for r_idx, row in enumerate(table.rows):
        if r_idx == 0: continue  # column headers
        cells = row.cells
        c0 = cells[0].text.strip() if cells else ''
        if not c0: continue  # empty row

        # Section header?
        sec_text = detect_section(row)
        if sec_text:
            current_section = {
                "id": SECTION_RE.match(sec_text).group(0),    # '1.0' / '2.0' / '3.0'
                "title": sec_text,
                "rows": [],
            }
            sections.append(current_section)
            continue

        if current_section is None:
            continue

        # Activity row — col 0 has like "1.1 Site Selection..."
        # The activity id is the leading "N.N" and the activity text is the rest.
        activity_match = ACTIVITY_RE.match(c0)
        activity_id   = activity_match.group(1) if activity_match else None
        activity_text = c0.replace('\n', ' ').strip()

        impacts     = cell_items(cells[1]) if len(cells) > 1 else []
        mitigations = cell_items(cells[2]) if len(cells) > 2 else []
        monitoring  = cell_items(cells[3]) if len(cells) > 3 else []
        person_impl = cell_plain(cells[4]) if len(cells) > 4 else None
        person_mon  = cell_plain(cells[5]) if len(cells) > 5 else None
        timeframe   = cell_plain(cells[6]) if len(cells) > 6 else None

        if not (impacts or mitigations or monitoring):
            continue

        row_id = f"{activity_id or f's{len(sections)}'}.r{len(current_section['rows'])+1}"
        current_section["rows"].append({
            "id": row_id,
            "activity_id": activity_id,
            "activity": activity_text,
            "impacts":     [{"id": f"{row_id}.i{i+1}", "label": x} for i, x in enumerate(impacts)],
            "mitigations": [{"id": f"{row_id}.m{i+1}", "label": x} for i, x in enumerate(mitigations)],
            "monitoring":  [{"id": f"{row_id}.p{i+1}", "label": x} for i, x in enumerate(monitoring)],
            "default_person_implement": person_impl,
            "default_person_monitor":   person_mon,
            "default_timeframe":        timeframe,
        })

    return {
        "title": title,
        "version": version,
        "kind": "emmp",
        "sections": sections,
    }

results = {}
for f in sorted(os.listdir(BASE)):
    if not f.endswith('.docx'): continue
    path = f"{BASE}/{f}"
    # is it actually local?
    try:
        with open(path, 'rb') as fh:
            chunk = fh.read(100)
            if len(chunk) < 100:
                print(f"  SKIP (stub):       {f}", file=sys.stderr)
                continue
    except OSError:
        print(f"  SKIP (read fail):  {f}", file=sys.stderr)
        continue
    if 'SCREENING' in f.upper() or 'COMPLIANCE' in f.upper():
        print(f"  TODO (later):      {f}", file=sys.stderr)
        continue
    try:
        schema = parse_emmp(path)
        if not schema:
            print(f"  EMPTY:             {f}", file=sys.stderr); continue
        total_rows  = sum(len(s["rows"]) for s in schema["sections"])
        total_items = sum(
            sum(len(r["impacts"])+len(r["mitigations"])+len(r["monitoring"]) for r in s["rows"])
            for s in schema["sections"]
        )
        sec_summary = " · ".join(f"{s['id']}={len(s['rows'])}r" for s in schema["sections"])
        print(f"  OK  {f}: {len(schema['sections'])} sections [{sec_summary}], {total_items} items", file=sys.stderr)
        results[f] = schema
    except Exception as e:
        print(f"  ERR {f}: {e}", file=sys.stderr)

with open(OUT, 'w') as fh:
    json.dump(results, fh, indent=2)
print(f"\nWrote {len(results)} schemas to {OUT}", file=sys.stderr)
